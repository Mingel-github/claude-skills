#!/usr/bin/env python3
"""
Generate publishable Word document from survey.md.

Features:
  - Word built-in heading styles (Heading 1/2/3) → TOC-ready
  - Auto TOC field (right-click → Update Field in Word)
  - Clickable cross-references: [N] in text → bookmark links to references
  - References: hanging indent + justified alignment + IEEE italic journal names
  - Fonts: SimSun (Chinese) + Times New Roman (English), no MS Gothic
  - Tables with shaded header rows
  - Inline formatting: **bold**, *italic*, $LaTeX$

Usage:
  python gen_docx.py <survey.md> <output.docx>
"""

import re, sys, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT_CN = '宋体'
FONT_EN = 'Times New Roman'

SIZES = {
    'title': 18,
    'h1': 14,
    'h2': 12,
    'h3': 11,
    'body': 11,
    'table': 8,
    'ref': 9,
}

# ---------------------------------------------------------------------------
# Low-level font/style helpers
# ---------------------------------------------------------------------------

def _set_style_fonts(style, cn=FONT_CN, en=FONT_EN):
    """Ensure a paragraph style uses SimSun + Times New Roman at XML level."""
    style.font.name = en
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:cs'), en)


def _set_run_font(run, size=SIZES['body'], bold=False, italic=False):
    """Set east-Asian + Latin fonts on a single run."""
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_EN
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rFonts.set(qn('w:cs'), FONT_EN)


def _add_run(para, text, size=SIZES['body'], bold=False, italic=False):
    """Append a formatted run to a paragraph."""
    run = para.add_run(text)
    _set_run_font(run, size=size, bold=bold, italic=italic)
    return run

# ---------------------------------------------------------------------------
# Hyperlink helper (internal bookmark links)
# ---------------------------------------------------------------------------

def _add_hyperlink(para, text, bookmark, size=SIZES['body']):
    """Add an internal hyperlink (clickable cross-reference) to a paragraph."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rPr.append(rFonts)
    # Size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))  # half-points
    rPr.append(sz)
    # Color (blue for hyperlink)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    para._element.append(hyperlink)


def _add_bookmark(para, bookmark_id):
    """Add a Word bookmark at the start of a paragraph."""
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), str(bookmark_id))
    bookmark_start.set(qn('w:name'), f'ref_{bookmark_id}')
    para._element.insert(0, bookmark_start)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), str(bookmark_id))
    para._element.append(bookmark_end)


# ---------------------------------------------------------------------------
# TOC insertion
# ---------------------------------------------------------------------------

def _add_toc(doc):
    """Insert a Word TOC field. User right-clicks → Update Field to populate."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = para.add_run()
    _set_run_font(run, size=SIZES['body'])

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar_sep)

    hint = para.add_run('（在 Word 中右键此处 → "更新域" 以生成目录）')
    _set_run_font(hint, size=9, italic=True)

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    hint._element.append(fldChar_end)

    doc.add_paragraph()  # blank line after TOC


# ---------------------------------------------------------------------------
# Paragraph builders
# ---------------------------------------------------------------------------

def _add_rich_para(doc, text, size=SIZES['body'], bold=False, italic=False,
                   indent=False, space_after=4, alignment=None, style=None):
    """
    Add a paragraph, parsing inline markup: **bold**, *italic*, $math$.
    Returns (paragraph, list of [N] citation numbers found).
    """
    p = doc.add_paragraph()
    if style:
        p.style = style
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.25
    if indent:
        pf.first_line_indent = Cm(0.74)
    if alignment is not None:
        p.alignment = alignment

    citations = []

    # Split on inline markup tokens
    segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\$\$[^$]+\$\$|\$[^$]+\$)', text)

    for seg in segments:
        if not seg:
            continue

        # Display math $$...$$
        if seg.startswith('$$') and seg.endswith('$$'):
            _add_run(p, seg[2:-2], size=size, italic=True)
        # Bold **...**
        elif seg.startswith('**') and seg.endswith('**'):
            _add_run(p, seg[2:-2], size=size, bold=True)
        # Italic *...* (single, not double)
        elif seg.startswith('*') and seg.endswith('*') and not seg.startswith('**'):
            _add_run(p, seg[1:-1], size=size, italic=True)
        # Math $...$ (but not $$)
        elif seg.startswith('$') and seg.endswith('$') and not seg.startswith('$$'):
            _add_run(p, seg[1:-1], size=size, italic=True)
        # Plain text — scan for [N] citations
        else:
            # Split plain text on [N] patterns to create hyperlinks
            sub_parts = re.split(r'(\[\d+(?:[,，\s]*\d+)*\])', seg)
            for sp in sub_parts:
                if not sp:
                    continue
                m = re.match(r'^\[(\d+(?:[,，\s]*\d+)*)\]$', sp)
                if m:
                    # Could be [1] or [1,2,3] or [1][2]
                    nums = re.findall(r'\d+', m.group(1))
                    for n in nums:
                        citations.append(int(n))
                    # Render as hyperlink
                    _add_hyperlink(p, sp, f'ref_{nums[0]}', size=size)
                else:
                    _add_run(p, sp, size=size)

    return p, citations


# ---------------------------------------------------------------------------
# Table parser
# ---------------------------------------------------------------------------

def _flush_table(doc, table_lines):
    """Convert accumulated markdown table lines into a Word table."""
    if not table_lines:
        return

    rows = []
    for tl in table_lines:
        # Protect | inside $...$ before splitting
        protected = re.sub(r'(\$[^$]+\$)', lambda m: m.group(1).replace('|', '\x00'), tl)
        cells = [c.strip().replace('\x00', '|') for c in protected.split('|')]
        cells = [c for c in cells if c]  # remove empty from leading/trailing |
        if cells and not all(re.match(r'^[-:\s]+$', c) for c in cells):
            rows.append(cells)

    if len(rows) < 2:
        return

    header = rows[0]
    body = [r for r in rows[1:] if len(r) == len(header)]
    if not body:
        return

    ncols = len(header)
    nrows = 1 + len(body)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, ct in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(ct)
        _set_run_font(run, size=SIZES['table'], bold=True)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._element.get_or_add_tcPr().append(shading)

    for i_row, row_data in enumerate(body):
        for j, ct in enumerate(row_data):
            cell = table.rows[i_row + 1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(ct)
            _set_run_font(run, size=SIZES['table'])

    doc.add_paragraph()  # spacer


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path, docx_path):
    doc = Document()

    # -- Configure built-in heading styles -----------------------------------
    for lvl, sz in [(1, SIZES['h1']), (2, SIZES['h2']), (3, SIZES['h3'])]:
        style = doc.styles[f'Heading {lvl}']
        style.font.size = Pt(sz)
        style.font.bold = True
        style.font.color.rgb = None  # auto black
        _set_style_fonts(style)
        style.paragraph_format.space_before = Pt(14 if lvl == 1 else 10)
        style.paragraph_format.space_after = Pt(6)

    # Normal style
    style_normal = doc.styles['Normal']
    style_normal.font.size = Pt(SIZES['body'])
    _set_style_fonts(style_normal)

    # Title style
    style_title = doc.styles['Title']
    style_title.font.size = Pt(SIZES['title'])
    style_title.font.bold = True
    _set_style_fonts(style_title)

    # -- Page margins ---------------------------------------------------------
    for sec in doc.sections:
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(3.18)
        sec.right_margin = Cm(3.18)

    # -- Read markdown --------------------------------------------------------
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    import os
    i = 0
    in_refs = False
    ref_bookmark_counter = 2000  # start high to avoid conflicts
    table_lines = []
    toc_added = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Empty line
        if not line:
            _flush_table(doc, table_lines)
            table_lines = []
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # H1 (# title)
        if line.startswith('# ') and not line.startswith('## '):
            _flush_table(doc, table_lines)
            table_lines = []
            p = doc.add_paragraph(line[2:].strip(), style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add TOC after title
            if not toc_added:
                doc.add_paragraph()
                _add_toc(doc)
                toc_added = True

            i += 1
            continue

        # H2 (## section)
        if line.startswith('## '):
            _flush_table(doc, table_lines)
            table_lines = []
            text = line[3:].strip()
            doc.add_paragraph(text, style='Heading 1')
            if text == '参考文献':
                in_refs = True
            i += 1
            continue

        # H3 (### subsection)
        if line.startswith('### '):
            _flush_table(doc, table_lines)
            table_lines = []
            doc.add_paragraph(line[4:].strip(), style='Heading 2')
            i += 1
            continue

        # Sub-sub-heading: standalone bold line (entirely **...**)
        if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            _flush_table(doc, table_lines)
            table_lines = []
            text = line.strip('*').strip()
            doc.add_paragraph(text, style='Heading 3')
            i += 1
            continue

        # Table rows
        if line.startswith('|') and line.endswith('|'):
            table_lines.append(line)
            i += 1
            continue

        # Flush pending table
        if table_lines:
            _flush_table(doc, table_lines)
            table_lines = []

        # --- Body paragraph --------------------------------------------------
        if in_refs:
            # Reference paragraph: hanging indent + justified
            p, _ = _add_rich_para(doc, line, size=SIZES['ref'], space_after=2,
                                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.74)
            pf.first_line_indent = Cm(-0.74)  # hanging indent

            # Add bookmark for cross-reference
            m = re.match(r'^\[(\d+)\]', line.strip())
            if m:
                ref_num = int(m.group(1))
                _add_bookmark(p, ref_num)
        else:
            p, _ = _add_rich_para(doc, line, indent=not line.startswith('-'))

        i += 1

    # Flush remaining
    _flush_table(doc, table_lines)

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python gen_docx.py <survey.md> <output.docx>')
        sys.exit(1)
    out = convert(sys.argv[1], sys.argv[2])
    print(f'Saved {out}')
